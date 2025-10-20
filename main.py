import os
import chardet
import tkinter as tk
from tkinter import filedialog, messagebox
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random
import string

NS = {'xs': 'http://www.w3.org/2001/XMLSchema'}


def resolve_path(schema_location, parent_file):
    parent_dir = os.path.dirname(parent_file)
    return os.path.normpath(os.path.join(parent_dir, schema_location))


def get_doc(node):
    if node is None:
        return ""
    doc = node.find('.//xs:annotation/xs:documentation[1]', namespaces=NS)
    if doc is None:
        doc = node.find('.//xs:annotation/xs:documentation', namespaces=NS)
    text_content = ""
    if doc is not None:
        text_content = (doc.text or "")
        for child in doc:
            text_content += etree.tostring(child, encoding='unicode', method='text') or ""
            text_content += (child.tail or "")
        text_content += (doc.tail or "")
    return text_content.strip().replace('\n', ' ').replace('\r', ' ')


class XSDDocumentationGenerator:
    def __init__(self):
        self.visited_files = set()
        self.schemas = {}
        self.enum_types = {}
        self.namespaces = {}
        self.simple_types = []
        self.root_elements = {}

    def load_schema(self, file_path):
        file_path = os.path.normpath(file_path)
        if file_path in self.visited_files:
            return
        self.visited_files.add(file_path)

        if not os.path.exists(file_path):
            print(f"Файл не найден: {file_path}")
            return

        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected['encoding']
            if encoding is None or encoding.lower() in ('ascii', 'utf-8', 'utf-8-sig'):
                encoding = 'utf-8'

            parser = etree.XMLParser(encoding=encoding, recover=True)
            tree = etree.parse(file_path, parser)
            root = tree.getroot()

            schema_name = os.path.basename(file_path)
            if schema_name.lower().endswith('.xsd'):
                schema_name = schema_name[:-4]
            schema_doc = get_doc(root)

            global_elements = root.xpath('/xs:schema/xs:element[@name]', namespaces=NS)

            first_element_name = ""
            first_element_doc = ""
            if global_elements:
                first_elem = global_elements[0]
                first_element_name = first_elem.get('name')
                first_element_doc = get_doc(first_elem)

            complex_types = {}
            for t in root.xpath('//xs:complexType[@name]', namespaces=NS):
                name = t.get('name')
                complex_types[name] = t

            for t in root.xpath('//xs:simpleType[@name]', namespaces=NS):
                name = t.get('name')
                enumerations = t.xpath('.//xs:enumeration', namespaces=NS)
                is_enum = len(enumerations) > 0

                base_type = ""
                restrictions = {}
                restriction_elem = t.find('.//xs:restriction', namespaces=NS)
                if restriction_elem is not None:
                    base_type = restriction_elem.get('base', '')
                    for facet in restriction_elem:
                        tag = str(facet.tag)
                        if tag.startswith(f"{{{NS['xs']}}}"):
                            facet_name = tag[len(f"{{{NS['xs']}}}"):]

                            if facet_name in [
                                'minLength', 'maxLength', 'minInclusive', 'maxInclusive',
                                'minExclusive', 'maxExclusive', 'pattern', 'totalDigits',
                                'fractionDigits'
                            ]:
                                restrictions[facet_name] = facet.get('value', '')

                type_record = {
                    'name': name,
                    'file': schema_name,
                    'base_type': base_type,
                    'restrictions': restrictions,
                    'description': get_doc(t),
                    'is_enum': is_enum
                }
                self.simple_types.append(type_record)

                if is_enum:
                    enum_values = []
                    for enum in enumerations:
                        code = enum.get('value')
                        desc = get_doc(enum)
                        if not desc:
                            desc = code
                        enum_values.append((code, desc))
                    self.enum_types[name] = enum_values

            for inc in root.xpath('.//xs:include | .//xs:import', namespaces=NS):
                loc = inc.get('schemaLocation')
                if loc:
                    next_path = resolve_path(loc, file_path)
                    self.load_schema(next_path)

            self.schemas[file_path] = {
                'name': schema_name,
                'doc': schema_doc,
                'first_element_name': first_element_name,
                'first_element_doc': first_element_doc,
                'complex_types': complex_types,
                'global_elements': global_elements,
                'root': root,
                'tree': tree
            }

            if global_elements:
                self.root_elements[schema_name] = global_elements[0]

            print(f"Загружена схема: {schema_name}")

        except Exception as e:
            print(f"Ошибка при загрузке {file_path}: {e}")
            import traceback
            traceback.print_exc()

    def add_shading(self, cell, color='D9D9D9'):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def create_table_with_header(self, doc, headers, widths=None):
        table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
        table.autofit = False

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.add_shading(hdr_cells[i])

        if widths:
            for i, w in enumerate(widths):
                if i < len(table.columns):
                    table.columns[i].width = w

        return table

    def add_row_to_table(self, table, values):
        cells = table.add_row().cells
        for i, text in enumerate(values):
            if i < len(cells):
                cell_text = str(text) if text is not None else ""
                cells[i].text = cell_text
                for p in cells[i].paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in p.runs:
                        run.font.size = Pt(10)
                cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def describe_type(self, table, node, schema_info, level=0):
        _ = schema_info
        if node.tag == f"{{{NS['xs']}}}complexType" and node.get('name'):
            type_name = node.get('name')
            self.add_row_to_table(table, [type_name, "", "блок", get_doc(node), ""])
            seq_all = node.xpath('.//xs:sequence | .//xs:all', namespaces=NS)
            if seq_all:
                for elem in seq_all[0].xpath('xs:element', namespaces=NS):
                    self.describe_element(table, elem, schema_info, level + 1)
            self.add_row_to_table(table, [f"/{type_name}", "", "", "Конец блока", ""])
            return

        if node.tag == f"{{{NS['xs']}}}complexType" and not node.get('name'):
            seq_all = node.xpath('.//xs:sequence | .//xs:all', namespaces=NS)
            if seq_all:
                for elem in seq_all[0].xpath('xs:element', namespaces=NS):
                    self.describe_element(table, elem, schema_info, level + 1)
            return

    def describe_element(self, table, elem, schema_info, level=0):
        _ = schema_info
        name = elem.get('name')
        ref = elem.get('ref')
        type_name = elem.get('type')
        min_occurs = elem.get('minOccurs', '1')
        required = "Да" if min_occurs == '1' else "Нет"

        if elem.getparent().tag == f"{{{NS['xs']}}}choice":
            required = "Нет\nВыбор"

        if ref:
            local_ref = ref.split(':')[-1] if ':' in ref else ref
            self.add_row_to_table(table, ["", name, f"ref: {local_ref}", get_doc(elem), required])
            return

        if type_name:
            local_type = type_name.split(':')[-1] if ':' in type_name else type_name
            if local_type in schema_info['complex_types']:
                self.add_row_to_table(table, ["", name, local_type, get_doc(elem), required])
            elif local_type in self.enum_types:
                self.add_row_to_table(table, ["", name, local_type, get_doc(elem), required])
            else:
                self.add_row_to_table(table, ["", name, local_type, get_doc(elem), required])
            return

        complex_type = elem.find('xs:complexType', namespaces=NS)
        if complex_type is not None:
            self.add_row_to_table(table, ["", name, "блок", get_doc(elem), ""])
            self.describe_type(table, complex_type, schema_info, level + 1)
            self.add_row_to_table(table, ["", f"/{name}", "", "Конец блока", ""])
            return

        self.add_row_to_table(table, ["", name, type_name or "string", get_doc(elem), required])

    def add_data_types_dictionary(self, doc):
        doc.add_heading('6. Словарь типов данных', level=1)

        if not self.simple_types:
            doc.add_paragraph("Простые типы данных не найдены.")
            return

        string_types = {}
        decimal_types = {}
        integer_types = {}
        datetime_types = {}
        other_types = {}

        for type_info in self.simple_types:
            type_name = type_info['name']
            base_type = type_info['base_type'].lower()
            if 'string' in base_type:
                string_types.setdefault(type_name, []).append(type_info)
            elif 'decimal' in base_type:
                decimal_types.setdefault(type_name, []).append(type_info)
            elif any(t in base_type for t in ['int', 'integer', 'long', 'short', 'byte']):
                integer_types.setdefault(type_name, []).append(type_info)
            elif any(t in base_type for t in ['date', 'time']):
                datetime_types.setdefault(type_name, []).append(type_info)
            else:
                other_types.setdefault(type_name, []).append(type_info)

        if string_types:
            doc.add_heading('6.1. Строковые типы (string)', level=2)
            headers = ["Имя XSD-файла", "Имя типа", "Базовый тип", "Мин. длина", "Макс. длина", "Паттерн", "Как заполняется"]
            widths = [Inches(1.2), Inches(1.0), Inches(0.8), Inches(0.6), Inches(0.6), Inches(1.0), Inches(2.0)]
            table = self.create_table_with_header(doc, headers, widths)
            for type_name in sorted(string_types.keys()):
                for info in string_types[type_name]:
                    restrictions = info['restrictions']
                    row = [
                        info['file'],
                        type_name,
                        info['base_type'],
                        restrictions.get('minLength', ''),
                        restrictions.get('maxLength', ''),
                        restrictions.get('pattern', ''),
                        info['description']
                    ]
                    self.add_row_to_table(table, row)

        if decimal_types:
            doc.add_heading('6.2. Десятичные типы (decimal)', level=2)
            headers = ["Имя XSD-файла", "Имя типа", "Базовый тип", "Мин. значение", "Макс. значение", "Всего цифр", "Дробных цифр", "Как заполняется"]
            widths = [Inches(1.0), Inches(1.0), Inches(0.8), Inches(0.7), Inches(0.7), Inches(0.6), Inches(0.6), Inches(1.8)]
            table = self.create_table_with_header(doc, headers, widths)
            for type_name in sorted(decimal_types.keys()):
                for info in decimal_types[type_name]:
                    restrictions = info['restrictions']
                    row = [
                        info['file'],
                        type_name,
                        info['base_type'],
                        restrictions.get('minInclusive', restrictions.get('minExclusive', '')),
                        restrictions.get('maxInclusive', restrictions.get('maxExclusive', '')),
                        restrictions.get('totalDigits', ''),
                        restrictions.get('fractionDigits', ''),
                        info['description']
                    ]
                    self.add_row_to_table(table, row)

        if integer_types:
            doc.add_heading('6.3. Целочисленные типы (int)', level=2)
            headers = ["Имя XSD-файла", "Имя типа", "Базовый тип", "Мин. значение", "Макс. значение", "Как заполняется"]
            widths = [Inches(1.2), Inches(1.0), Inches(0.8), Inches(0.8), Inches(0.8), Inches(2.2)]
            table = self.create_table_with_header(doc, headers, widths)
            for type_name in sorted(integer_types.keys()):
                for info in integer_types[type_name]:
                    restrictions = info['restrictions']
                    row = [
                        info['file'],
                        type_name,
                        info['base_type'],
                        restrictions.get('minInclusive', restrictions.get('minExclusive', '')),
                        restrictions.get('maxInclusive', restrictions.get('maxExclusive', '')),
                        info['description']
                    ]
                    self.add_row_to_table(table, row)

        if datetime_types:
            doc.add_heading('6.4. Типы даты и времени (date и dateTime)', level=2)
            headers = ["Имя XSD-файла", "Имя типа", "Базовый тип", "Мин. значение", "Макс. значение", "Шаблон (формат)", "Как заполняется"]
            widths = [Inches(1.0), Inches(1.0), Inches(0.8), Inches(0.8), Inches(0.8), Inches(1.0), Inches(1.8)]
            table = self.create_table_with_header(doc, headers, widths)
            for type_name in sorted(datetime_types.keys()):
                for info in datetime_types[type_name]:
                    restrictions = info['restrictions']
                    row = [
                        info['file'],
                        type_name,
                        info['base_type'],
                        restrictions.get('minInclusive', restrictions.get('minExclusive', '')),
                        restrictions.get('maxInclusive', restrictions.get('maxExclusive', '')),
                        restrictions.get('pattern', ''),
                        info['description']
                    ]
                    self.add_row_to_table(table, row)

        if other_types:
            doc.add_heading('6.5. Остальные типы', level=2)
            headers = ["Имя XSD-файла", "Имя типа", "Базовый тип", "Как заполняется"]
            widths = [Inches(1.5), Inches(1.5), Inches(1.0), Inches(3.2)]
            table = self.create_table_with_header(doc, headers, widths)
            for type_name in sorted(other_types.keys()):
                for info in other_types[type_name]:
                    row = [
                        info['file'],
                        type_name,
                        info['base_type'],
                        info['description']
                    ]
                    self.add_row_to_table(table, row)

    def generate_sample_value(self, type_name, schema_info):
        if type_name in self.enum_types and self.enum_types[type_name]:
            return self.enum_types[type_name][0][0]

        for type_info in self.simple_types:
            if type_info['name'] == type_name:
                base_type = type_info['base_type'].lower()
                restrictions = type_info['restrictions']

                if 'string' in base_type:
                    min_len = int(restrictions.get('minLength', 1))
                    max_len = int(restrictions.get('maxLength', max(min_len, 10)))
                    length = min(max_len, max(min_len, 5))
                    pattern = restrictions.get('pattern', '')
                    if pattern:
                        if 'digit' in pattern.lower() or '\\d' in pattern:
                            return ''.join(random.choices(string.digits, k=length))
                        elif '[a-z]' in pattern.lower():
                            return ''.join(random.choices(string.ascii_lowercase, k=length))
                        elif '[a-zA-Z0-9]' in pattern or 'alnum' in pattern.lower():
                            return ''.join(random.choices(string.ascii_letters + string.digits, k=length))
                    return ''.join(random.choices(string.ascii_letters, k=length))

                elif any(t in base_type for t in ['int', 'integer', 'long', 'short', 'byte']):
                    min_val = int(
                        float(restrictions.get('minInclusive', restrictions.get('minExclusive', -2147483648))))
                    max_val = int(float(restrictions.get('maxInclusive', restrictions.get('maxExclusive', 2147483647))))
                    min_val = max(min_val, -2147483648)
                    max_val = min(max_val, 2147483647)
                    return str(random.randint(min_val, max_val))

                elif 'decimal' in base_type or 'double' in base_type or 'float' in base_type:
                    min_val = float(restrictions.get('minInclusive', restrictions.get('minExclusive', -1e6)))
                    max_val = float(restrictions.get('maxInclusive', restrictions.get('maxExclusive', 1e6)))
                    fraction_digits = int(restrictions.get('fractionDigits', 2))
                    value = round(random.uniform(min_val, max_val), fraction_digits)
                    return f"{value:.{fraction_digits}f}"

                elif 'date' in base_type and 'time' not in base_type:
                    return "2025-10-20"

                elif 'datetime' in base_type or ('date' in base_type and 'time' in base_type):
                    return "2025-10-20T12:00:00"

                elif 'boolean' in base_type:
                    return "true"

        if 'string' in type_name.lower():
            return "Пример текста"
        elif any(t in type_name.lower() for t in ['int', 'integer']):
            return "123"
        elif 'decimal' in type_name.lower() or 'double' in type_name.lower() or 'float' in type_name.lower():
            return "123.45"
        elif 'date' in type_name.lower() and 'time' not in type_name.lower():
            return "2025-10-20"
        elif 'time' in type_name.lower():
            return "12:00:00"
        elif 'boolean' in type_name.lower():
            return "true"

        return "Пример значения"

    def find_complex_type(self, type_name):
        for schema_info in self.schemas.values():
            if type_name in schema_info['complex_types']:
                return schema_info['complex_types'][type_name], schema_info
        return None, None

    def generate_xml_example(self, element, schema_info, level=0):
        indent = "  " * level
        name = element.get('name')
        if not name:
            return []

        type_name = element.get('type')
        min_occurs = element.get('minOccurs', '1')
        max_occurs = element.get('maxOccurs', '1')

        if min_occurs == '0':
            return []

        result = []

        if type_name:
            local_type = type_name.split(':')[-1] if ':' in type_name else type_name

            ct, ct_schema = self.find_complex_type(local_type)
            if ct is not None:
                attr_parts = []
                all_attributes = []

                def collect_attributes(node, node_schema):
                    attrs = node.xpath('.//xs:attribute', namespaces=NS)
                    for attr in attrs:
                        aname = attr.get('name')
                        if not aname:
                            continue
                        use = attr.get('use', 'optional')
                        atype = attr.get('type', 'string')
                        local_atype = atype.split(':')[-1] if ':' in atype else atype
                        avalue = self.generate_sample_value(local_atype, node_schema or schema_info)
                        all_attributes.append((aname, avalue, use))

                    extension = node.find('.//xs:extension', namespaces=NS)
                    if extension is not None:
                        base_ref = extension.get('base')
                        if base_ref:
                            base_local = base_ref.split(':')[-1] if ':' in base_ref else base_ref
                            base_ct, base_schema = self.find_complex_type(base_local)
                            if base_ct is not None:
                                collect_attributes(base_ct, base_schema)

                collect_attributes(ct, ct_schema or schema_info)

                attr_str = ""
                if all_attributes:
                    attr_str = " " + " ".join(f'{aname}="{avalue}"' for aname, avalue, _ in all_attributes)

                result.append(f"{indent}<{name}{attr_str}>")

                all_elements = []

                def collect_elements(node, node_schema):
                    extension = node.find('.//xs:extension', namespaces=NS)
                    if extension is not None:
                        base_ref = extension.get('base')
                        if base_ref:
                            base_local = base_ref.split(':')[-1] if ':' in base_ref else base_ref
                            base_ct, base_schema = self.find_complex_type(base_local)
                            if base_ct is not None:
                                collect_elements(base_ct, base_schema)

                        ext_seq = extension.xpath('xs:sequence | xs:all', namespaces=NS)
                        for seq in ext_seq:
                            for elem in seq.xpath('xs:element', namespaces=NS):
                                all_elements.append((elem, node_schema))
                    else:
                        seqs = node.xpath('.//xs:sequence | .//xs:all', namespaces=NS)
                        for seq in seqs:
                            for elem in seq.xpath('xs:element', namespaces=NS):
                                all_elements.append((elem, node_schema))

                collect_elements(ct, ct_schema or schema_info)

                for child_elem, child_schema in all_elements:
                    child_min = child_elem.get('minOccurs', '1')
                    if child_min != '0':
                        result.extend(self.generate_xml_example(child_elem, child_schema, level + 1))

                result.append(f"{indent}</{name}>")
            else:
                value = self.generate_sample_value(local_type, schema_info)
                result.append(f"{indent}<{name}>{value}</{name}>")
        else:
            complex_type = element.find('xs:complexType', namespaces=NS)
            if complex_type is not None:
                attr_parts = []
                for attr in complex_type.xpath('.//xs:attribute', namespaces=NS):
                    aname = attr.get('name')
                    if not aname:
                        continue
                    atype = attr.get('type', 'string')
                    local_atype = atype.split(':')[-1] if ':' in atype else atype
                    avalue = self.generate_sample_value(local_atype, schema_info)
                    attr_parts.append(f'{aname}="{avalue}"')
                attr_str = " " + " ".join(attr_parts) if attr_parts else ""

                result.append(f"{indent}<{name}{attr_str}>")

                seqs = complex_type.xpath('.//xs:sequence | .//xs:all', namespaces=NS)
                for seq in seqs:
                    for child_elem in seq.xpath('xs:element', namespaces=NS):
                        child_min = child_elem.get('minOccurs', '1')
                        if child_min != '0':
                            result.extend(self.generate_xml_example(child_elem, schema_info, level + 1))

                result.append(f"{indent}</{name}>")
            else:
                result.append(f"{indent}<{name}>Пример значения</{name}>")

        return result

    def add_xml_examples(self, doc):
        doc.add_heading('7. Примеры XML файлов', level=1)

        if not self.root_elements:
            doc.add_paragraph("Корневые элементы не найдены.")
            return

        for schema_name, root_element in self.root_elements.items():
            doc.add_heading(f'7.1. Пример XML для схемы {schema_name}', level=2)

            schema_info = None
            for path, info in self.schemas.items():
                if info['name'] == schema_name:
                    schema_info = info
                    break

            if schema_info:
                try:
                    example_lines = self.generate_xml_example(root_element, schema_info, 0)
                    example_xml = '\n'.join(example_lines)
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(example_xml)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                except Exception as e:
                    doc.add_paragraph(f"Ошибка при генерации примера: {str(e)}")
                    import traceback
                    traceback.print_exc()
            else:
                doc.add_paragraph("Не удалось сгенерировать пример: информация о схеме не найдена.")

    def generate_docx(self, xsd_paths, output_path):
        for path in xsd_paths:
            self.load_schema(path)

        doc = Document()

        title = doc.add_heading('Описание форматов электронных документов', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()

        doc.add_heading('Содержание:', level=1)
        content_items = [
            "1. Термины, определения и сокращения",
            "2. Общие положения",
            "3. Перечень электронных документов",
            "4. Справочник XML-структур"
        ]

        for i, schema in enumerate(sorted(self.schemas.values(), key=lambda x: x['name']), start=1):
            content_items.append(f"4.{i}. {schema['name']}")
        content_items.append("5. Справочник глобальных кодов")

        for i, etype in enumerate(sorted(self.enum_types.keys()), start=1):
            content_items.append(f"5.{i}. {etype}")

        content_items.append("6. Словарь типов данных")

        if any('string' in info['base_type'].lower() for info in self.simple_types):
            content_items.append("6.1. Строковые типы (string)")
        if any('decimal' in info['base_type'].lower() for info in self.simple_types):
            content_items.append("6.2. Десятичные типы (decimal)")
        if any(any(t in info['base_type'].lower() for t in ['int', 'integer', 'long', 'short', 'byte']) for info in self.simple_types):
            content_items.append("6.3. Целочисленные типы (int)")
        if any(any(t in info['base_type'].lower() for t in ['date', 'time']) for info in self.simple_types):
            content_items.append("6.4. Типы даты и времени (date и dateTime)")
        content_items.append("6.5. Остальные типы")
        content_items.append("7. Примеры XML файлов")

        for i, schema_name in enumerate(sorted(self.root_elements.keys()), start=1):
            content_items.append(f"7.{i}. Пример XML для схемы {schema_name}")

        for item in content_items:
            doc.add_paragraph(item, style='List Number')
        doc.add_page_break()

        doc.add_heading('1. Термины, определения и сокращения', level=1)
        doc.add_paragraph('XML – Extensible Markup Language, расширяемый язык разметки.')
        doc.add_page_break()

        doc.add_heading('2. Общие положения', level=1)
        doc.add_paragraph('Общие положения отсутствуют.')
        doc.add_page_break()

        doc.add_heading('3. Перечень электронных документов', level=1)
        msg_table = doc.add_table(rows=1, cols=2, style='Table Grid')
        hdr_cells = msg_table.rows[0].cells
        for i, h in enumerate(["Наименование документа", "ID документа <XMLMsgNm>"]):
            p = hdr_cells[i].paragraphs[0]
            p.add_run(h).bold = True
            self.add_shading(hdr_cells[i])
        for schema in sorted(self.schemas.values(), key=lambda x: x['name']):
            doc_name = schema['first_element_doc'] if schema['first_element_doc'] else "—"
            msg_id = schema['name'] if schema['name'] else "—"
            self.add_row_to_table(msg_table, [doc_name, msg_id])
        doc.add_page_break()

        doc.add_heading('4. Справочник XML-структур', level=1)
        for i, (path, schema) in enumerate(sorted(self.schemas.items()), start=1):
            doc.add_heading(f"4.{i}. {schema['name']}", level=2)
            table = self.create_table_with_header(doc, [
                "Имя XML-типа в словаре",
                "Название XML-элемента в блоке",
                "Тип данных",
                "Содержание и значение XML-элемента",
                "Обязательность в XML-типе"
            ], [Inches(1.8), Inches(1.4), Inches(1.3), Inches(2.5), Inches(1.0)])

            types_used_by_elements = set()
            if schema['global_elements']:
                for elem in schema['global_elements']:
                    elem_type = elem.get('type')
                    if elem_type:
                        local_type_name = elem_type.split(':')[-1] if ':' in elem_type else elem_type
                        if local_type_name in schema['complex_types']:
                            types_used_by_elements.add(local_type_name)

            if schema['global_elements']:
                for elem in schema['global_elements']:
                    name = elem.get('name')
                    type_name = elem.get('type')
                    min_occurs = elem.get('minOccurs', '1')
                    required = "Да" if min_occurs == '1' else "Нет"

                    if elem.getparent().tag == f"{{{NS['xs']}}}choice":
                        required = "Нет\nВыбор"

                    if type_name:
                        local_type = type_name.split(':')[-1] if ':' in type_name else type_name
                        self.add_row_to_table(table, ["", name, local_type, get_doc(elem), required])
                        if local_type in schema['complex_types']:
                            self.describe_type(table, schema['complex_types'][local_type], schema, 1)
                    else:
                        complex_type = elem.find('xs:complexType', namespaces=NS)
                        if complex_type is not None:
                            self.add_row_to_table(table, ["", name, "блок", get_doc(elem), required])
                            self.describe_type(table, complex_type, schema, 1)
                            self.add_row_to_table(table, ["", f"/{name}", "", "Конец блока", ""])
                        else:
                            self.add_row_to_table(table, ["", name, "string", get_doc(elem), required])
            else:
                doc.add_paragraph("Глобальные элементы не найдены.")

            standalone_types = set(schema['complex_types'].keys()) - types_used_by_elements
            if standalone_types:
                doc.add_paragraph("Автономные типы (не связанные напрямую с элементами):")
                for type_name in sorted(standalone_types):
                    ct = schema['complex_types'][type_name]
                    self.describe_type(table, ct, schema, 0)

            doc.add_paragraph()
        doc.add_page_break()

        doc.add_heading('5. Справочник глобальных кодов', level=1)
        enum_entries = [st for st in self.simple_types if st.get('is_enum', False) and st['name'] in self.enum_types]

        if not enum_entries:
            doc.add_paragraph("Перечисления (enum) не найдены.")
        else:
            sorted_enum_entries = sorted(enum_entries, key=lambda x: (x['file'], x['name']))
            for i, st in enumerate(sorted_enum_entries, start=1):
                type_name = st['name']
                file_name = st['file']
                description = st['description']
                values = self.enum_types[type_name]

                doc.add_heading(f"5.{i}. {type_name}", level=2)
                if description.strip():
                    p = doc.add_paragraph()
                    p.add_run(description).italic = True

                enum_table = doc.add_table(rows=1, cols=3, style='Table Grid')
                hdr_cells = enum_table.rows[0].cells
                headers = ["Имя XSD-файла", "Код", "Описание"]
                for col_idx, header in enumerate(headers):
                    p = hdr_cells[col_idx].paragraphs[0]
                    p.add_run(header).bold = True
                    self.add_shading(hdr_cells[col_idx])

                for code, desc in values:
                    row_cells = enum_table.add_row().cells
                    row_cells[0].text = file_name
                    row_cells[1].text = code
                    row_cells[2].text = desc
                    for cell in row_cells:
                        for p in cell.paragraphs:
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            for run in p.runs:
                                run.font.size = Pt(10)
                doc.add_paragraph()
        doc.add_page_break()

        self.add_data_types_dictionary(doc)
        doc.add_page_break()

        self.add_xml_examples(doc)

        doc.save(output_path)


def main():
    root = tk.Tk()
    root.withdraw()

    messagebox.showinfo("Выбор XSD", "Выберите все XSD-файлы схемы")
    xsd_paths = filedialog.askopenfilenames(
        title="Выберите XSD-файлы",
        filetypes=[("XSD файлы", "*.xsd")],
        multiple=True
    )
    if not xsd_paths:
        messagebox.showwarning("Отмена", "Выбор файлов отменен.")
        return

    docx_path = filedialog.asksaveasfilename(
        title="Сохранить DOCX",
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )
    if not docx_path:
        messagebox.showwarning("Отмена", "Сохранение отменено.")
        return

    try:
        gen = XSDDocumentationGenerator()
        gen.generate_docx(xsd_paths, docx_path)
        messagebox.showinfo("Готово", f"Документация сохранена:\n{docx_path}")
    except Exception as e:
        import traceback
        error_msg = f"Ошибка: {str(e)}\n\n{traceback.format_exc()}"
        print(error_msg)
        messagebox.showerror("Ошибка", f"Произошла ошибка при генерации документа:\n{str(e)}")


if __name__ == "__main__":
    main()